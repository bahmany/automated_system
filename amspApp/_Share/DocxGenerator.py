import os
from html.parser import HTMLParser
import uuid
# from ams.settings import ABSOLUTE_PATH
# from dms.classes.word_tst import convertHtml2docx
# from docx import *

import uuid

from bs4 import BeautifulSoup, NavigableString, Tag
from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import CT_RPr
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from amsp import settings
from amspApp.FileServer.views.FileUploadView import FileUploadViewSet
from amspApp.Infrustructures.Classes.DateConvertors import mil_to_sh


def PrepareTables(tbl):
    finishedTB = True
    endPosOfTB = 0
    tbls = []
    tables = []
    finalTables = []
    while finishedTB:
        startPosOfTB = tbl.find("<table", endPosOfTB)
        if startPosOfTB == -1:
            finishedTB = False
            break
        startPosOfTB_closed = tbl.find(">", startPosOfTB + 1)
        endPosOfTB = tbl.find("</table>", startPosOfTB_closed + 1)
        endTd = startPosOfTB_closed
        tables.append([startPosOfTB, startPosOfTB_closed + 1, endPosOfTB])

        finishedTR = True
        endPosOfTr = startPosOfTB_closed
        positions = []
        while finishedTR:
            startPosOfTr = tbl.find("<tr", endPosOfTr)
            if startPosOfTr >= endPosOfTB or startPosOfTr == -1:
                finishedTR = False
                break
            startPosOfTr_closed = tbl.find(">", startPosOfTr + 1)
            endPosOfTr = tbl.find("</tr>", startPosOfTr_closed + 1)
            endTd = startPosOfTr_closed
            finishedTD = True
            tds = []
            while finishedTD:
                startTd = tbl.find("<td", endTd + 1)
                startTdClose = tbl.find(">", startTd + 1)
                endTd = tbl.find("</td>", startTdClose + 1)
                if endTd >= endPosOfTr or startTd == -1:
                    finishedTD = False
                    break
                tds.append([startTd, startTdClose, endTd])
            positions.append([startPosOfTr, startPosOfTr_closed, endPosOfTr, tds])
        newPos = []
        for i in positions:
            newTr = []
            newTr = [i[0], i[1], i[2], [b for b in i[3][::-1]]]
            newPos.append(newTr)

        finalStr = ""
        for i in newPos:
            td_str = ""
            for b in i[3]:
                td_str = td_str + tbl[b[0]:b[2]] + "</td>"
            finalStr = finalStr + tbl[i[0]:i[1] + 1] + td_str + "</tr>"
        finalStr = tbl[startPosOfTB:startPosOfTB_closed + 1] + finalStr + "</table>"
        finalStr = finalStr
        finalTables.append([startPosOfTB, startPosOfTB_closed, endPosOfTB, finalStr])

    for i in finalTables[::-1]:
        tbl = tbl.replace(tbl[i[0]:i[2]], i[3])

    tbl = tbl.replace("</table></table>", "</table>")
    return tbl


class MLStripper(HTMLParser):
    def __init__(self):
        self.reset()
        self.fed = []

    def handle_data(self, d):
        self.fed.append(d)

    def get_data(self):
        return ''.join(self.fed)


def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()


htmlstr = u"""<p dir="rtl"><strong>??????&zwnj;??? ???? ????? ?? ????? ? ?????? ????</strong></p>

<p dir="rtl">?- ????? <strong>????</strong> <em>???</em> ??? <strong>???????</strong></p>

<p dir="rtl">?- ????? <u>????</u> ?????</p>

<p dir="rtl">?- ???? ?????? <em>???????</em> ? ???? ?? ????</p>

<table border="1" cellpadding="1" cellspacing="1" style="width:500px">
	<tbody>
		<tr>
			<td>???</td>
			<td>????</td>
		</tr>
		<tr>
			<td><strong>???? ???</strong></td>
			<td>?</td>
		</tr>
		<tr>
			<td>??? ???? ?????</td>
			<td>?</td>
		</tr>
	</tbody>
</table>

<p dir="rtl">?- ???? ????&zwnj;??? ?????? ? ???? ?? ????</p>


"""


def perpareHtmlText(txt):
    # txt = unicode(txt)
    txt = txt.replace("\n", "")
    txt = txt.replace("\r", "")
    txt = txt.replace("\l", "")
    txt = txt.replace("<strong>", " #112 ")
    txt = txt.replace("</strong>", " #113 ")
    txt = txt.replace("<em>", " #115 ")
    txt = txt.replace("</em>", " #116 ")
    txt = txt.replace("<u>", " #118 ")
    txt = txt.replace("</u>", " #119 ")
    return txt


def makeWithStyle(newPar, bs4htmltext, thisIsTable=False):
    txt = bs4htmltext.__str__()
    txt = perpareHtmlText(txt)
    txt = strip_tags(txt)
    isItBold = False
    isItItalic = False
    isItUnderLine = False
    ttxt = txt.split(" ")
    for nP in ttxt:
        if nP == "#112": isItBold = True
        if nP == "#115": isItItalic = True
        if nP == "#118": isItUnderLine = True
        if nP == "#113": isItBold = False
        if nP == "#116": isItItalic = False
        if nP == "#119": isItUnderLine = False
        if len(nP) > 0:
            if nP[0] == "#": continue
        if nP == '': continue
        # nP = perpareHtmlText(nP)
        run = newPar.add_run(nP + ' ')
        if isItBold: run.bold = True
        if isItItalic: run.italic = True
        if isItUnderLine: run.underline = True
        # newPar.add_run(u" ")


def convertHtml2docx(
        God="",
        Shoar="",
        Recievers="",
        Numbering={},
        htmlstr="",
        Sign="",
        RooNevesht="",
        baseTemplate="",
        specificDate=None
):
    htmlstr = PrepareTables(htmlstr)
    document = Document(os.path.abspath(baseTemplate))

    if God != "":
        GodParag = document.add_paragraph(God, style='GodStyle')

    if Shoar != "":
        ShoarParag = document.add_paragraph(Shoar, style='ShoarStyle')

    if Numbering:
        NumberingParag = document.add_paragraph(style='NumberingStyle')

        Nrun = NumberingParag.add_run()
        Nrun.add_picture(Numbering["numbering"], height=Pt(10))
        Nrun.add_picture(Numbering["numberinglbl"], height=Pt(10))
        NumberingParag.add_run("\n")
        Drun = NumberingParag.add_run()
        Drun.add_picture(Numbering["date"], height=Pt(10))
        Drun.add_picture(Numbering["datelbl"], height=Pt(10))
        NumberingParag.add_run("\n")
        Prun = NumberingParag.add_run()
        Prun.add_picture(Numbering["peivast"], height=Pt(10))
        Prun.add_picture(Numbering["peivastlbl"], height=Pt(10))
        NumberingParag.add_run("\n")

    document.add_paragraph("", "Normal")

    if Recievers != "":
        document.add_paragraph("", "Normal")
        RecieversParag = document.add_paragraph(Recievers, style='RecieversStyle')

    document.add_paragraph("", "Normal")
    soup = BeautifulSoup(htmlstr, 'html.parser')
    currentObj = soup.p
    objs = []
    while currentObj != None:
        objs.append(currentObj)
        currentObj = currentObj.next_element
    for i in soup.contents:
        newPar = None
        if type(i) == NavigableString:
            continue
        if type(i) == Tag:
            if i.name == "p":
                newPar = document.add_paragraph(style='Normal')
                newPar.style.font.rtl = True
                makeWithStyle(newPar, i)
        if i.name == "table":
            rowsCount = i.find_all("tr").__len__()
            for ii in range(0, rowsCount):
                colsCount = i.find_all("tr")[ii].find_all("td").__len__()
            table = document.add_table(rowsCount, colsCount)
            table.style = 'TableGrid'
            table.style.font.rtl = True
            for ii in range(0, rowsCount):
                currentColCount = len(i.find_all("tr")[ii].find_all("td"))
                for yy in range(0, currentColCount):
                    cell = table.cell(ii, yy)
                    cellParag = cell.paragraphs[0]
                    try :
                        makeWithStyle(
                            cellParag,
                            i.find_all("tr")[ii].find_all("td")[yy],
                            thisIsTable=True)
                    except:
                        cell = cell

    document.add_paragraph("", "Normal")
    #
    # document.add_paragraph("","Normal")
    #
    if Sign != "":
        SignParag = document.add_paragraph(Sign, style='SignStyle')

    if RooNevesht != "":
        # document.add_paragraph("","Normal")
        document.add_paragraph("", "Normal")
        RooNeveshtParag = document.add_paragraph(RooNevesht, style='RooneveshtStyle')
    fileName = uuid.uuid4().hex
    fileNameAndPath = FileUploadViewSet().getFolderName(specificDate, fileName)
    document.save(fileNameAndPath)
    return {"filename": fileName, "fullPath": fileNameAndPath}


class MLStripper(HTMLParser):
    def __init__(self):
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.fed = []

    def handle_data(self, d):
        self.fed.append(d)

    def get_data(self):
        return ''.join(self.fed)


def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()


def makePng(outPut):
    import PIL.Image as I, PIL.ImageFont as IF, PIL.ImageDraw as ID

    def change_to_persian(outPut):
        replace = [("1", u"۱"), ("2", u"۲"), ("3", u"۳"), ("4", u"۴"), ("5", u"۵"), ("6", u"۶"), ("7", u"۷"),
                   ("۸", u"?"),
                   ("۹", u"?"), ("۰", u"?")]
        outChar = ""
        cc = ""
        for o in outPut:
            outChar = o
            for rr in replace:
                if o == rr[0]:
                    o = rr[1]
                    outChar = rr[1]
            cc = cc + outChar
        return cc

    output_pics = []
    totalWidth = 0
    v = 0
    fontsize = 35
    # f = IF.truetype(ABSOLUTE_PATH.func_globals["BASE_DIR"] + "/dms/static/css/fonts/Yekan.ttf", fontsize, )
    f = IF.truetype(settings.APP_PATH + "static/fonts/bnazanin.ttf", fontsize, )

    for a in outPut:
        v = v + 1
        t = a
        i = I.new("L", (1000, 1000), color=255)
        d = ID.Draw(i)

        d.text((0, 0), t, fill=0, font=f)
        text_width, text_height = d.textsize(t, font=f)
        # calcing offset
        i.size = (text_width, text_height + (fontsize - text_height))
        totalWidth = totalWidth + i.size[0]
        output_pics.append(i)
        # i.save("/tmp/du_" + str(v) + ".png", optimize=1)
    i = I.new("L", (totalWidth + 10, i.size[1]), color=255)
    d = ID.Draw(i)
    x = -1
    xx = 0
    from bidi.algorithm import get_display

    for a in outPut:
        x = x + 1
        t = get_display(change_to_persian(a), upper_is_rtl=True)
        # t = a
        sizeOf = (xx, 0)
        d.text(sizeOf, t, fill=0, font=f)
        xx = xx + output_pics[x].size[0]
    filename = uuid.uuid1().hex + ".png"

    filename = FileUploadViewSet().getFolderName(datetime.now(), filename)
    i.save(filename)

    return {"outPut": outPut, "fileAddr": filename}


# def LetterTextToDownload(itemsToReplace):
# from ams import settings
#
# projectPath = settings.PROJECT_ROOT
#
# templ_newpath = projectPath + "dms/static/templates/****A4.docx"
# if itemsToReplace["typeOf"] == "a4": templ_newpath = projectPath + "dms/static/templates/****A4.docx"
#     if itemsToReplace["typeOf"] == "a5": templ_newpath = projectPath + "dms/static/templates/****A5.docx"
#     doc = Document(docx=templ_newpath)
#     style = doc.styles['Normal']
#     font = style.font
#     font.name = 'B Nazanin'
#     font.size = Pt(10)
#     """
#     253151 : God
#     55895  : Somareh
#     98547  : Tarikh
#     58942  : Peivast
#     889845 : LetterTo
#     5955554: Body
#     558548 : EmzaKonandeh
#     778587 : Roonevesht
#     """
#
#     for i in itemsToReplace.keys():
#         itemsToReplace[i] = itemsToReplace[i].replace("&zwnj;", " ")
#         itemsToReplace[i] = itemsToReplace[i].replace("&laquo;", "«")
#         itemsToReplace[i] = itemsToReplace[i].replace("&raquo;", "»")
#
#     for par in doc.paragraphs:
#         for i in itemsToReplace.keys():
#             if par.text.find(unicode(i)) > -1:
#                 par.text = par.text.replace(unicode(i), strip_tags(itemsToReplace[i]))
#                 par.style.font.name = "B Nazanin"
#                 par.style.font.rtl = True
#
#     for tbl in doc.tables:
#         for cell in tbl._cells:
#             for i in itemsToReplace.keys():
#                 for par in cell.paragraphs:
#                     if par.text.find(unicode(i)) > -1:
#
#                         if par.text.find('55895') > -1:
#                             if itemsToReplace['55895'] != "":
#                                 par.text = par.text.replace(u'55895', u'')
#                                 par.runs[0].add_picture(
#                                     itemsToReplace['55895'],
#                                     height=Pt(12)
#                                 )
#                         else:
#                             par.text = par.text.replace(unicode(i), strip_tags(itemsToReplace[i]))
#
#                         par.style.font.name = "B Nazanin"
#                         par.style.font.rtl = True
#
#     pathOfNewFile = projectPath + 'dms/files/'
#     fileName = "prev_" + uuid.uuid4().hex + '.docx'
#
#     doc.save(pathOfNewFile + fileName)
#     return "/filesprev/" + fileName
def LetterTextToDownload(itemsToReplace, templateAddr):
    """
    253151 : Header
    55895  : Somareh img address
    111599  : Sign img address

    98547  : Tarikh
    58942  : Peivast
    889845 : LetterTo
    5955554: Body
    558548 : Sign Text
    778587 : Footer Text >> roonevesht recievers puts here
    """
    if not "111599" in itemsToReplace:
        itemsToReplace["111599"] = ""
    if itemsToReplace["111599"] == "":
        itemsToReplace["111599"] = "22/22/22"

    if itemsToReplace["58942"] == "":
        itemsToReplace["58942"] = "0"
    if itemsToReplace["98547"] == "":
        itemsToReplace["98547"] = mil_to_sh(datetime.now().strftime("%Y-%m-%d"))
    # dateImg = makePng("/".join(list(reversed(itemsToReplace["98547"].split("/")))))
    dateImg = makePng("/".join(list((itemsToReplace["98547"].split("/")))))
    # shomarehImg = makePng("/".join(list(reversed(itemsToReplace["55895"].split("/")))))
    shomarehImg = makePng(itemsToReplace["111599"])
    peiImg = makePng(itemsToReplace["58942"])
    somareh = {
        "numbering": shomarehImg["fileAddr"],
        "numberinglbl": settings.APP_PATH + "static/images/Shomareh.png",

        "date": dateImg["fileAddr"],
        "datelbl": settings.APP_PATH + "static/images/Tarikh.png",

        "peivast": peiImg["fileAddr"],
        "peivastlbl": settings.APP_PATH + "static/images/Peivast.png",
    }

    # if itemsToReplace["typeOf"] == u'A4':

    # if itemsToReplace["typeOf"] == u'A5':
    #     path = settings.APP_PATH + "static/templates/" + "****A5.docx"
    #
    return convertHtml2docx(
        God=itemsToReplace["253153"],
        Shoar="",
        Recievers=itemsToReplace["889845"],
        Numbering=somareh,
        htmlstr=itemsToReplace["5955554"].replace("div", "p"),
        RooNevesht=itemsToReplace["778587"],
        Sign=itemsToReplace["558548"],
        baseTemplate=templateAddr,
        specificDate=itemsToReplace["dateToDownload"]
    )
